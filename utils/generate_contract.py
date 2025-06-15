print("✅ generate_contract.py is being read!")

from docx import Document
from datetime import datetime
import os
import sys
import uuid
import shutil
from docx.oxml import parse_xml

# Add the project root to the Python path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from models import db
from models.provider import Provider
from models.standard_rates import StandardRates  # We'll need to create this model

# Standard Imaging Rates (replace or load from DB as needed)
IMAGING_RATES = {
    "MRI w/o": 300,
    "MRI w/": 400,
    "MRI w/ & w/o": 450,
    "CT w/o": 200,
    "CT w/": 275,
    "CT w/ & w/o": 350,
    "XRAY": 25,
    "Arthrograms": 570
}

# Imaging categories for selection
IMAGING_CATEGORIES = [
    "MRI w/o",
    "MRI w/",
    "MRI w/ & w/o",
    "CT w/o",
    "CT w/",
    "CT w/ & w/o",
    "XRAY",
    "Arthrograms"
]

def get_rates_by_method(provider_id, method, custom_rates=None, wcfs_percentages=None):
    """
    Get rates based on the selected reimbursement method.
    
    Args:
        provider_id: The ID of the provider
        method: The reimbursement method ('standard', 'wcfs', or 'custom')
        custom_rates: Dictionary of custom rates for each category (for 'custom' method)
        wcfs_percentages: Dictionary of WCFS percentages for each category (for 'wcfs' method)
    
    Returns:
        Dictionary of rates for each imaging category and state
    """
    provider = Provider.query.get(provider_id)
    if not provider:
        raise ValueError("Provider not found")
    
    # Get states from provider
    if provider.states_in_contract:
        states = [s.strip() for s in provider.states_in_contract.split(",") if s.strip()]
    else:
        states = ["CA"]  # Default to California if no states specified
    
    rates = {}
    
    # Map form field names to category names
    category_mapping = {
        'mri_without': 'MRI w/o',
        'mri_with': 'MRI w/',
        'mri_both': 'MRI w/ & w/o',
        'ct_without': 'CT w/o',
        'ct_with': 'CT w/',
        'ct_both': 'CT w/ & w/o',
        'xray': 'XRAY',
        'arthrogram': 'Arthrograms'
    }
    
    if method == 'standard':
        # Get rates from standard_rates table for each state and category
        for state in states:
            rates[state] = {}
            for category in IMAGING_CATEGORIES:
                standard_rate = StandardRates.query.filter_by(
                    state=state,
                    category=category
                ).first()
                
                if standard_rate:
                    rates[state][category] = standard_rate.rate
                else:
                    # Fallback to default rates if not found in DB
                    rates[state][category] = IMAGING_RATES.get(category, 0)
    
    elif method == 'wcfs':
        # Apply WCFS percentages to standard rates
        if not wcfs_percentages:
            raise ValueError("WCFS percentages must be provided for WCFS method")
        
        # Store the WCFS percentages in the rates dictionary
        for state in states:
            rates[state] = {}
            for form_field, category in category_mapping.items():
                if form_field in wcfs_percentages.get(state, {}):
                    # Store the percentage as an integer
                    rates[state][category] = int(wcfs_percentages[state][form_field])
                else:
                    # Use 0 if no percentage provided
                    rates[state][category] = 0
    
    elif method == 'custom':
        # Use custom rates provided by the user
        if not custom_rates:
            raise ValueError("Custom rates must be provided for custom method")
        
        for state in states:
            rates[state] = {}
            for form_field, category in category_mapping.items():
                if form_field in custom_rates.get(state, {}):
                    rates[state][category] = custom_rates[state][form_field]
                else:
                    # Use default rate if no custom rate provided
                    rates[state][category] = IMAGING_RATES.get(category, 0)
    
    else:
        raise ValueError(f"Invalid reimbursement method: {method}")
    
    return rates

def create_rates_table(doc, states, rates, method='standard', wcfs_percentages=None):
    # Create table without relying on named styles
    table = doc.add_table(rows=1, cols=3)
    
    # Make sure the table has borders using direct XML
    for row in table.rows:
        for cell in row.cells:
            # Add borders to the cell
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Add borders using direct XML
            for key in ['top', 'left', 'bottom', 'right']:
                element = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:{key} w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>')
                tcPr.append(element)
            
            # Make header text bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "State"
    hdr_cells[1].text = "Procedure"
    hdr_cells[2].text = "Rate" if method != 'wcfs' else "WCFS %"

    for state in states:
        for procedure in IMAGING_CATEGORIES:
            row_cells = table.add_row().cells
            
            # Add borders to each new cell
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                
                # Add borders using direct XML
                for key in ['top', 'left', 'bottom', 'right']:
                    element = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:{key} w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>')
                    tcPr.append(element)
            
            row_cells[0].text = state
            row_cells[1].text = procedure
            
            if method == 'wcfs':
                # Get the percentage for this state and procedure
                if wcfs_percentages and state in wcfs_percentages:
                    percentage = wcfs_percentages[state].get(procedure, 0)
                else:
                    percentage = 0
                row_cells[2].text = f"{percentage}% of WCFS"
            else:
                # Get the rate for this state and procedure
                if state in rates and procedure in rates[state]:
                    rate = rates[state][procedure]
                else:
                    rate = 0
                row_cells[2].text = f"${rate:.2f}"
            
    return table

def generate_contract(provider_id, method='standard', custom_rates=None, wcfs_percentages=None):
    provider = Provider.query.get(provider_id)
    if not provider:
        raise ValueError("Provider not found")

    print(f"🔍 Generating contract for provider: {provider.name}")
    print(f"🔍 Method: {method}")
    print(f"🔍 Provider data: name={provider.name}, dba={provider.dba_name}, address={provider.address}")

    # Split state list - handle None case
    if provider.states_in_contract:
        states = [s.strip() for s in provider.states_in_contract.split(",") if s.strip()]
    else:
        states = ["CA"]  # Default to California if no states specified

    print(f"🔍 States: {states}")

    # Get rates based on the selected method
    rates = get_rates_by_method(provider_id, method, custom_rates, wcfs_percentages)
    print(f"🔍 Rates: {rates}")

    # Load template - ensure we use the exact path specified
    template_path = "templates/contracts/IMAGING_TEMPLATE.docx"
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")

    print(f"🔍 Template found at: {template_path}")

    # Create output paths
    output_folder = "contracts"
    os.makedirs(output_folder, exist_ok=True)
    
    # Create filename using DBA name or provider name if DBA is not available
    filename = f"{provider.dba_name or provider.name}_Agreement"
    # Replace any characters that might be invalid in filenames
    filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).strip()
    
    docx_path = os.path.join(output_folder, f"{filename}.docx")
    pdf_path = os.path.join(output_folder, f"{filename}.pdf")

    print(f"🔍 Output paths - DOCX: {docx_path}, PDF: {pdf_path}")

    # Create a copy of the template for this contract
    shutil.copy2(template_path, docx_path)
    
    # Open the copy for editing
    doc = Document(docx_path)

    # Import XML parser for table formatting
    from docx.oxml import parse_xml
    
    print("🔍 Starting placeholder replacement...")
    
    # Replace all placeholders in paragraphs
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text
        if "{{provider_name}}" in text:
            print(f"🔍 Replacing provider_name with: {provider.name}")
            text = text.replace("{{provider_name}}", provider.name)
            paragraph.text = text
        elif "{{dba_name}}" in text:
            print(f"🔍 Replacing dba_name with: {provider.dba_name or ''}")
            text = text.replace("{{dba_name}}", provider.dba_name or "")
            paragraph.text = text
        elif "{{address}}" in text:
            print(f"🔍 Replacing address with: {provider.address or ''}")
            text = text.replace("{{address}}", provider.address or "")
            paragraph.text = text
        elif "{{provider_type}}" in text:
            print(f"🔍 Replacing provider_type with: {provider.provider_type or ''}")
            text = text.replace("{{provider_type}}", provider.provider_type or "")
            paragraph.text = text
        elif "{{npi}}" in text:
            print(f"🔍 Replacing npi with: {provider.npi or ''}")
            text = text.replace("{{npi}}", provider.npi or "")
            paragraph.text = text
        elif "{{specialty}}" in text:
            print(f"🔍 Replacing specialty with: {provider.specialty or ''}")
            text = text.replace("{{specialty}}", provider.specialty or "")
            paragraph.text = text
        elif "{{rate_type}}" in text:
            print(f"🔍 Replacing rate_type with: {method}")
            text = text.replace("{{rate_type}}", method)
            paragraph.text = text
        elif "{{wcfs_percentage}}" in text:
            if method == 'wcfs' and wcfs_percentages:
                # For WCFS method, show the average percentage
                avg_percentage = sum(wcfs_percentages.values()) / len(wcfs_percentages)
                print(f"🔍 Replacing wcfs_percentage with: {avg_percentage:.1f}")
                text = text.replace("{{wcfs_percentage}}", f"{avg_percentage:.1f}")
            else:
                print("🔍 Replacing wcfs_percentage with: 0")
                text = text.replace("{{wcfs_percentage}}", "0")
            paragraph.text = text
        elif "{{states}}" in text:
            states_text = ", ".join(states)
            print(f"🔍 Replacing states with: {states_text}")
            text = text.replace("{{states}}", states_text)
            paragraph.text = text
        elif "{{date}}" in text:
            current_date = datetime.now().strftime("%B %d, %Y")
            print(f"🔍 Replacing date with: {current_date}")
            text = text.replace("{{date}}", current_date)
            paragraph.text = text
        elif "{{exhibit_a}}" in text:
            print("🔍 Creating rates table...")
            # Create the rates table
            table = create_rates_table(doc, states, rates, method, wcfs_percentages)
            
            # Insert the table after the current paragraph
            p = paragraph._element
            p.getparent().insert(p.getparent().index(p) + 1, table._element)
            
            # Remove the placeholder paragraph
            p.getparent().remove(p)
            
            # Add a small paragraph after the table for spacing
            doc.add_paragraph()
            print("🔍 Rates table created and inserted")

    print("🔍 Placeholder replacement completed")

    # Save the modified copy
    doc.save(docx_path)
    print(f"🔍 Document saved to: {docx_path}")

    # Convert to PDF using LibreOffice in headless mode
    pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        # Use LibreOffice in headless mode to convert DOCX to PDF
        convert_cmd = f'soffice --headless --convert-to pdf --outdir "{os.path.dirname(pdf_path)}" "{docx_path}"'
        result = os.system(convert_cmd)
        
        if result == 0 and os.path.exists(pdf_path):
            print(f"✅ PDF generated using LibreOffice: {pdf_path}")
        else:
            print(f"ℹ️ PDF generation failed, only DOCX available")
            pdf_path = None
                
    except Exception as e:
        print(f"🔍 [WARN] PDF conversion failed: {e}")
        pdf_path = None

    return docx_path, pdf_path

def generate_contract_per_state(provider_id, state_configurations):
    """
    Generate contract with different rate configurations per state.
    
    Args:
        provider_id: The ID of the provider
        state_configurations: Dict with state as key and config dict as value
                              e.g., {'CA': {'method': 'wcfs', 'wcfs_percentages': {...}}}
    
    Returns:
        Tuple of (docx_path, pdf_path)
    """
    provider = Provider.query.get(provider_id)
    if not provider:
        raise ValueError("Provider not found")

    # Parse states
    if provider.states_in_contract:
        states = [s.strip() for s in provider.states_in_contract.split(",") if s.strip()]
    else:
        states = ["CA"]

    # Load template
    template_path = "templates/contracts/IMAGING_TEMPLATE.docx"
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found at {template_path}")

    # Create output paths
    output_folder = "contracts"
    os.makedirs(output_folder, exist_ok=True)
    
    filename = f"{provider.dba_name or provider.name}_Agreement"
    filename = "".join(c for c in filename if c.isalnum() or c in (' ', '-', '_')).strip()
    
    docx_path = os.path.join(output_folder, f"{filename}.docx")
    pdf_path = os.path.join(output_folder, f"{filename}.pdf")

    # Create a copy of the template
    shutil.copy2(template_path, docx_path)
    doc = Document(docx_path)

    # Replace placeholders (same as before)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if "{{provider_name}}" in text:
            paragraph.text = text.replace("{{provider_name}}", provider.name or "")
        elif "{{dba_name}}" in text:
            paragraph.text = text.replace("{{dba_name}}", provider.dba_name or "")
        elif "{{address}}" in text:
            paragraph.text = text.replace("{{address}}", provider.address or "")
        elif "{{provider_type}}" in text:
            paragraph.text = text.replace("{{provider_type}}", provider.provider_type or "")
        elif "{{npi}}" in text:
            paragraph.text = text.replace("{{npi}}", provider.npi or "")
        elif "{{specialty}}" in text:
            paragraph.text = text.replace("{{specialty}}", provider.specialty or "")
        elif "{{states}}" in text:
            paragraph.text = text.replace("{{states}}", ", ".join(states))
        elif "{{date}}" in text:
            paragraph.text = text.replace("{{date}}", datetime.now().strftime("%B %d, %Y"))
        elif "{{exhibit_a}}" in text:
            # Create the per-state rates table
            table = create_per_state_rates_table(doc, states, state_configurations)
            
            # Insert the table
            p = paragraph._element
            p.getparent().insert(p.getparent().index(p) + 1, table._element)
            p.getparent().remove(p)
            doc.add_paragraph()

    doc.save(docx_path)

    # Convert to PDF
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
    except Exception as e:
        print(f"[WARN] PDF conversion failed: {e}")

    return docx_path, pdf_path

def create_per_state_rates_table(doc, states, state_configurations):
    """Create a rates table that shows different rate types per state."""
    from docx.oxml import parse_xml
    
    table = doc.add_table(rows=1, cols=3)
    
    # Add borders to table
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for key in ['top', 'left', 'bottom', 'right']:
                element = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:{key} w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>')
                tcPr.append(element)

    # Header
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "State"
    hdr_cells[1].text = "Procedure"
    hdr_cells[2].text = "Rate/Percentage"

    for state in states:
        config = state_configurations.get(state, {'method': 'standard'})
        method = config['method']
        
        # Get rates for this state based on its configuration
        if method == 'standard':
            # Get standard rates for this state
            state_rates = {}
            for category in IMAGING_CATEGORIES:
                standard_rate = StandardRates.query.filter_by(
                    state=state, category=category
                ).first()
                if standard_rate:
                    state_rates[category] = standard_rate.rate
                else:
                    state_rates[category] = IMAGING_RATES.get(category, 0)
        
        elif method == 'wcfs':
            state_rates = config.get('wcfs_percentages', {})
        
        elif method == 'custom':
            state_rates = config.get('custom_rates', {})
        
        for category in IMAGING_CATEGORIES:
            row_cells = table.add_row().cells
            
            # Add borders to new cells
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for key in ['top', 'left', 'bottom', 'right']:
                    element = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:{key} w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>')
                    tcPr.append(element)
            
            row_cells[0].text = state
            row_cells[1].text = category
            
            rate_value = state_rates.get(category, 0)
            if method == 'wcfs':
                row_cells[2].text = f"{rate_value}% of WCFS"
            else:
                row_cells[2].text = f"${rate_value:.2f}"
    
    return table

# Test section - can be run directly
if __name__ == "__main__":
    with app.app_context():
        # Create a test provider
        provider = Provider(
            id=str(uuid.uuid4()),
            name="Test Imaging Center",
            dba_name="TIC Medical Group",
            address="123 Test St, Los Angeles, CA 90001",
            provider_type="Imaging",
            states_in_contract="CA, NY, TX",
            npi="1234567890",
            specialty="Radiology",
            rate_type="wcfs",
            wcfs_percentage=80.0,
            status="active"
        )
        
        # Add to database
        db.session.add(provider)
        db.session.commit()
        
        print(f"Created test provider with ID: {provider.id}")
        
        # Generate contract
        try:
            docx_path, pdf_path = generate_contract(provider.id)
            print(f"Contract generated successfully!")
            print(f"DOCX path: {docx_path}")
            print(f"PDF path: {pdf_path}")
        except Exception as e:
            print(f"Error generating contract: {e}")