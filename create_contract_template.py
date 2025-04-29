from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_contract_template():
    doc = Document()
    
    # Title
    title = doc.add_heading('Provider Agreement', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Date
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run('{{ date }}')
    
    # Introduction
    doc.add_paragraph('This Provider Agreement (the "Agreement") is entered into as of {{ date }} by and between Clarity Diagnostics, Inc. ("Clarity") and {{ provider.name }} (the "Provider").')
    
    # Provider Information
    doc.add_heading('Provider Information', level=1)
    provider_info = doc.add_paragraph()
    provider_info.add_run('Provider Name: ').bold = True
    provider_info.add_run('{{ provider.name }}\n')
    provider_info.add_run('DBA Name: ').bold = True
    provider_info.add_run('{{ provider.dba_name }}\n')
    provider_info.add_run('Address: ').bold = True
    provider_info.add_run('{{ provider.address }}\n')
    provider_info.add_run('Provider Type: ').bold = True
    provider_info.add_run('{{ provider.provider_type }}\n')
    provider_info.add_run('NPI: ').bold = True
    provider_info.add_run('{{ provider.npi }}\n')
    provider_info.add_run('Specialty: ').bold = True
    provider_info.add_run('{{ provider.specialty }}')
    
    # Rate Structure
    doc.add_heading('Rate Structure', level=1)
    rate_info = doc.add_paragraph()
    rate_info.add_run('Rate Type: ').bold = True
    rate_info.add_run('{{ provider.rate_type }}\n')
    rate_info.add_run('{% if provider.rate_type == "wcfs" %}WCFS Percentage: {{ provider.wcfs_percentage }}%{% endif %}')
    
    # States
    doc.add_heading('States', level=1)
    states_para = doc.add_paragraph()
    states_para.add_run('This agreement applies to the following states: ')
    states_para.add_run('{% for state in states %}{{ state }}{% if not loop.last %}, {% endif %}{% endfor %}')
    
    # Exhibit A
    doc.add_heading('Exhibit A - Rate Schedule', level=1)
    doc.add_paragraph('The following rates apply to the services provided by the Provider:')
    
    # Placeholder for the table that will be generated dynamically
    doc.add_paragraph('{{ exhibit_a }}')
    
    # Save the template
    doc.save('templates/contracts/contract_imaging.docx')
    print("Contract template created successfully!")

if __name__ == "__main__":
    create_contract_template() 